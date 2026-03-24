import os
from dotenv import load_dotenv
load_dotenv()  # load .env FIRST

# LangSmith tracing — MUST use LANGCHAIN_ prefix, not LANGSMITH_
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"]    = os.getenv("LANGSMITH_API_KEY", "")
os.environ["LANGCHAIN_PROJECT"]    = "KnowBot"
os.environ["LANGCHAIN_ENDPOINT"]   = "https://api.smith.langchain.com"

import json
import shutil
from datetime import datetime
from typing import Optional
from fastapi import FastAPI, UploadFile, File, Query, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from rag_chain import load_rag_chain, get_answer
from sarvam_transcribe import transcribe_chunk, transcribe_file, generate_meeting_notes
from ingest import ingest_project, delete_project_vectors, delete_file_vectors
from storage import upload_file_to_s3, delete_file_from_s3, delete_project_from_s3

import os
from fastapi.responses import StreamingResponse
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import io
import httpx

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=False,
)

PROJECTS_FILE  = "projects.json"
PROJECTS_DIR   = "projects"
ALLOWED_EXTS   = {".pdf", ".pptx", ".xlsx"}

# In-memory cache for RAG chains per project
rag_cache = {}
meeting_transcripts = {}

def load_projects():
    if not os.path.exists(PROJECTS_FILE):
        return {"projects": []}
    with open(PROJECTS_FILE, "r") as f:
        return json.load(f)

def save_projects(data):
    with open(PROJECTS_FILE, "w") as f:
        json.dump(data, f, indent=2)

def get_rag(project_id: str = None):
    """Get or create RAG chain for a project."""
    key = project_id or "all"
    if key not in rag_cache:
        rag_cache[key] = load_rag_chain(project_id)
    return rag_cache[key]

def invalidate_rag_cache(project_id: str = None):
    """Clear cache when docs change."""
    key = project_id or "all"
    if key in rag_cache:
        del rag_cache[key]
    if "all" in rag_cache:
        del rag_cache["all"]

# ── Models ──
class QuestionRequest(BaseModel):
    question: str
    project_id: Optional[str] = None
    session_id: str = "default"

class CreateProjectRequest(BaseModel):
    name: str
    description: str = ""

class RenameProjectRequest(BaseModel):
    name: str

class EndMeetingRequest(BaseModel):
    session_id: str
    meeting_type: str = "internal"

# ══════════════════════════════════════
# ── Document Q&A ──
# ══════════════════════════════════════
@app.post("/ask")
async def ask_question(req: QuestionRequest):
    chain_tuple = get_rag(req.project_id)
    answer, sources, _ = get_answer(
        chain_tuple,
        req.question,
        session_id=req.session_id,
        project_id=req.project_id
    )
    return {"answer": answer, "sources": sources}

# ══════════════════════════════════════
# ── Project Management ──
# ══════════════════════════════════════
@app.get("/projects")
async def list_projects():
    data = load_projects()
    return data

@app.post("/projects")
async def create_project(req: CreateProjectRequest):
    data = load_projects()

    # Generate unique ID from name
    project_id = req.name.lower().strip().replace(" ", "-").replace("/", "-")
    project_id = ''.join(c for c in project_id if c.isalnum() or c == '-')

    # Check if already exists
    existing_ids = [p["id"] for p in data["projects"]]
    if project_id in existing_ids:
        raise HTTPException(status_code=400, detail="Project with this name already exists.")

    # Create folders
    docs_folder = os.path.join(PROJECTS_DIR, project_id, "docs")
    os.makedirs(docs_folder, exist_ok=True)

    # Add to projects.json
    new_project = {
        "id": project_id,
        "name": req.name,
        "description": req.description,
        "created_at": datetime.now().strftime("%Y-%m-%d"),
        "doc_count": 0,
        "docs": []
    }
    data["projects"].append(new_project)
    save_projects(data)

    print(f"Created project: {req.name} ({project_id})")
    return new_project

@app.put("/projects/{project_id}")
async def rename_project(project_id: str, req: RenameProjectRequest):
    data = load_projects()
    for project in data["projects"]:
        if project["id"] == project_id:
            project["name"] = req.name
            save_projects(data)
            invalidate_rag_cache(project_id)
            return project
    raise HTTPException(status_code=404, detail="Project not found.")

@app.delete("/projects/{project_id}")
async def delete_project(project_id: str):
    data = load_projects()
    project = next((p for p in data["projects"] if p["id"] == project_id), None)

    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    # 1. Delete vectors from Qdrant
    try:
        delete_project_vectors(project_id)
        print(f"[QDRANT] Deleted vectors for project: {project_id}")
    except Exception as e:
        print(f"[QDRANT] Delete warning: {e}")

    # 2. Delete all files from S3
    try:
        delete_project_from_s3(project_id)
        print(f"[S3] Deleted all files for project: {project_id}")
    except Exception as e:
        print(f"[S3] Delete warning: {e}")

    # 3. Delete local folder
    folder = os.path.join(PROJECTS_DIR, project_id)
    if os.path.exists(folder):
        shutil.rmtree(folder)
        print(f"[LOCAL] Deleted folder: {folder}")

    # 4. Remove from projects.json
    data["projects"] = [p for p in data["projects"] if p["id"] != project_id]
    save_projects(data)
    invalidate_rag_cache(project_id)

    return {"message": f"Project '{project_id}' deleted successfully."}


@app.delete("/projects/{project_id}/docs/{filename}")
async def delete_document(project_id: str, filename: str):
    data = load_projects()
    project = next((p for p in data["projects"] if p["id"] == project_id), None)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    # 1. Delete from S3
    try:
        delete_file_from_s3(project_id, filename)
        print(f"[S3] Deleted: {filename}")
    except Exception as e:
        print(f"[S3] Delete warning: {e}")

    # 2. Delete from local disk
    file_path = os.path.join(PROJECTS_DIR, project_id, "docs", filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"[LOCAL] Deleted: {file_path}")

    # 3. Delete vectors from Qdrant
    try:
        delete_file_vectors(project_id, filename)
        print(f"[QDRANT] Deleted vectors for: {filename}")
    except Exception as e:
        print(f"[QDRANT] Delete warning: {e}")

    # 4. Update projects.json
    project["docs"] = [d for d in project["docs"] if d != filename]
    project["doc_count"] = len(project["docs"])
    save_projects(data)
    invalidate_rag_cache(project_id)

    return {"message": f"{filename} deleted from project {project_id}."}




@app.post("/projects/{project_id}/upload")
async def upload_to_project(
    project_id: str,
    file: UploadFile = File(...)
):
    data = load_projects()
    project = next((p for p in data["projects"] if p["id"] == project_id), None)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"File format '{ext}' is not supported. Allowed: PDF, PPTX, XLSX"
        )

    # Save temporarily
    os.makedirs("temp", exist_ok=True)
    temp_path = os.path.join("temp", file.filename)
    file_bytes = await file.read()
    with open(temp_path, "wb") as f:
        f.write(file_bytes)

    # Upload to S3 — permanent cloud storage
    upload_file_to_s3(temp_path, project_id, file.filename)

    # Also save locally for ingestion into Qdrant
    docs_folder = os.path.join(PROJECTS_DIR, project_id, "docs")
    os.makedirs(docs_folder, exist_ok=True)
    local_path = os.path.join(docs_folder, file.filename)
    with open(local_path, "wb") as f:
        f.write(file_bytes)

    # Clean up temp file
    if os.path.exists(temp_path):
        os.remove(temp_path)

    print(f"Saved: {file.filename} to project {project_id}")

    # Ingest into Qdrant
    success = ingest_project(project_id, project["name"])

    if success:
        if file.filename not in project["docs"]:
            project["docs"].append(file.filename)
            project["doc_count"] = len(project["docs"])
        save_projects(data)
        invalidate_rag_cache(project_id)
        return {
            "message": f"{file.filename} uploaded and indexed successfully.",
            "project": project
        }
    else:
        raise HTTPException(status_code=500, detail="Indexing failed.")

# ══════════════════════════════════════
# ── Meeting Intelligence ──
# ══════════════════════════════════════
@app.post("/meeting/transcribe-chunk")
async def transcribe_meeting_chunk(
    session_id: str = Query(...),
    language: str = Query("en-IN"),
    file: UploadFile = File(...)
):
    audio_bytes = await file.read()
    print(f"[CHUNK] Session: {session_id} | Size: {len(audio_bytes)} bytes")

    chunk_transcript = transcribe_chunk(audio_bytes, language)
    print(f"[CHUNK] Transcript: '{chunk_transcript}'")

    if session_id not in meeting_transcripts:
        meeting_transcripts[session_id] = []

    if chunk_transcript.strip():
        meeting_transcripts[session_id].append(chunk_transcript)

    return {
        "chunk_transcript": chunk_transcript,
        "full_transcript": " ".join(meeting_transcripts.get(session_id, []))
    }

@app.post("/meeting/end")
async def end_meeting(req: EndMeetingRequest):
    session_id = req.session_id
    if session_id not in meeting_transcripts:
        return {"error": "No transcript found for this session."}

    full_transcript = " ".join(meeting_transcripts[session_id])
    if not full_transcript.strip():
        return {"error": "No speech detected during the meeting."}

    notes = generate_meeting_notes(full_transcript, req.meeting_type)
    notes["full_transcript"] = full_transcript
    del meeting_transcripts[session_id]
    return notes

@app.post("/meeting/transcribe-file")
async def transcribe_meeting_file(
    language: str = Query("en-IN"),
    meeting_type: str = Query("internal"),
    file: UploadFile = File(...)
):
    os.makedirs("temp", exist_ok=True)
    path = f"temp/{file.filename}"
    with open(path, "wb") as f:
        f.write(await file.read())

    transcript = transcribe_file(path, language)
    notes = generate_meeting_notes(transcript, meeting_type)
    notes["full_transcript"] = transcript
    return notes

@app.delete("/chat/session/{session_id}")
async def clear_chat_session(session_id: str):
    from memory import chat_memory
    chat_memory.clear_session(session_id)
    return {"message": f"Session {session_id} cleared."}

@app.post("/meeting/download/docx")
async def download_docx(data: dict):
    summary = data.get("summary", "")
    key_decisions = data.get("key_decisions", "")
    action_items = data.get("action_items", "")
    transcript = data.get("transcript", "")

    doc = Document()
    doc.add_heading("KnowBot — Meeting Summary", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary if summary else "No summary available.")
    doc.add_heading("Key Decisions", level=1)
    doc.add_paragraph(key_decisions if key_decisions else "No key decisions recorded.")
    doc.add_heading("Action Items", level=1)
    doc.add_paragraph(action_items if action_items else "No action items recorded.")
    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph(transcript if transcript else "No transcript available.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=meeting-summary.docx"}
    )


@app.post("/meeting/download/pdf")
async def download_pdf(data: dict):
    summary = data.get("summary", "")
    key_decisions = data.get("key_decisions", "")
    action_items = data.get("action_items", "")
    transcript = data.get("transcript", "")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("KnowBot — Meeting Summary", styles['Title']))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Summary", styles['Heading1']))
    story.append(Paragraph(summary if summary else "No summary available.", styles['Normal']))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Key Decisions", styles['Heading1']))
    story.append(Paragraph(key_decisions if key_decisions else "No key decisions recorded.", styles['Normal']))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Action Items", styles['Heading1']))
    story.append(Paragraph(action_items if action_items else "No action items recorded.", styles['Normal']))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Full Transcript", styles['Heading1']))
    story.append(Paragraph(transcript if transcript else "No transcript available.", styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=meeting-summary.pdf"}
    )

@app.post("/projects/{project_id}/upload-url")
async def upload_from_url(project_id: str, data: dict):
    url = data.get("url", "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL is required.")

    # Handle Google Drive links automatically
    if "drive.google.com/file/d/" in url:
        file_id = url.split("/file/d/")[1].split("/")[0]
        url = f"https://drive.google.com/uc?export=download&confirm=t&id={file_id}"
        filename = f"gdrive_{file_id}.pdf"
        ext = ".pdf"
    else:
        clean_url = url.split("?")[0]
        filename = clean_url.split("/")[-1]
        ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Only PDF, PPTX, XLSX allowed.")

    # Find project
    data_projects = load_projects()
    project = next((p for p in data_projects["projects"] if p["id"] == project_id), None)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    # Download file
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=60) as client:
            response = await client.get(url, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            })
            response.raise_for_status()
            content = response.content
            # Check if Google returned HTML instead of file
            if content[:5] in [b"<!DOC", b"<html", b"\xef\xbb\xbf<"]:
                raise HTTPException(status_code=400, detail="Google Drive returned a preview page. Make sure sharing is set to 'Anyone with the link'.")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to download file: {str(e)}")

    # Save to project docs folder
    docs_folder = os.path.join(PROJECTS_DIR, project_id, "docs")
    os.makedirs(docs_folder, exist_ok=True)
    file_path = os.path.join(docs_folder, filename)

    with open(file_path, "wb") as f:
        f.write(content)

    # Upload to S3
    try:
        upload_file_to_s3(file_path, project_id, filename)
    except Exception as e:
        print(f"[S3] Upload warning: {e}")

    # Ingest into Qdrant
    success = ingest_project(project_id, project["name"])
    if not success:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail="Indexing failed.")

    # Update projects.json
    if filename not in project["docs"]:
        project["docs"].append(filename)
        project["doc_count"] = len(project["docs"])
    save_projects(data_projects)
    invalidate_rag_cache(project_id)

    return {"project": project}
